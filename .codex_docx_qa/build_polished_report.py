import io
import os
import zipfile
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


SOURCE = Path(os.environ["MILKBOT_SOURCE_DOCX"])
OUTPUT = Path(os.environ["MILKBOT_OUTPUT_DOCX"])
ASSET_DIR = Path(os.environ["MILKBOT_ASSET_DIR"])

BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EEF5FA"
GRAY = "666666"
LIGHT_GRAY = "E7E6E6"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=110, bottom=100, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LIGHT_GRAY, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_run_font(run, size=None, bold=None, color=None, name="맑은 고딕"):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_text(paragraph, text, size=10.5, bold=False, color="222222"):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return run


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(17)
    section.bottom_margin = Mm(16)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)

    normal = doc.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.2


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(hp, "학회 참가 결과보고서", size=8.5, color=GRAY)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(fp, "박준서  |  학회 참가 및 증빙자료", size=8, color="7F7F7F")


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(4)
    add_text(p, "학회 참가 결과보고서", size=25, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(15)
    add_text(p, "참가자 박준서", size=12, color=GRAY)


def add_section_heading(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_with_next = True
    add_text(p, title, size=15, bold=True, color=BLUE)
    if subtitle:
        add_text(p, f"  {subtitle}", size=9, color=GRAY)
    return p


def add_info_table(doc):
    rows = [
        ("참가자", "박준서", "참가 기간", "8월 5일(수) ~ 8월 7일(금)"),
        ("이동 경로", "한국에너지공과대학교 → 광주버스터미널 → 경주", "주요 활동", "강연, 포스터 세션, Plenary Talk 참석"),
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(2.4), Cm(6.3), Cm(2.4), Cm(6.3)]
    for row, values in zip(table.rows, rows):
        for i, (cell, value) in enumerate(zip(row.cells, values)):
            cell.width = widths[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=115, bottom=115)
            if i % 2 == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i % 2 == 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_text(p, value, size=9.2, bold=(i % 2 == 0), color=BLUE if i % 2 == 0 else "222222")
    set_table_borders(table, color="B4C7D7", size=7)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_callout(doc, label, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(17.2)
    set_cell_shading(cell, PALE_BLUE)
    set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_text(p, f"{label}  ", size=10, bold=True, color=BLUE)
    add_text(p, body, size=10, color="333333")
    set_table_borders(table, color="B4C7D7", size=6)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def extract_assets():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(SOURCE) as zf:
        for i in range(1, 23):
            names = [n for n in zf.namelist() if n.startswith("word/media/") and Path(n).stem == f"image{i}"]
            if not names:
                raise ValueError(f"Missing image{i}")
            name = names[0]
            data = zf.read(name)
            image = Image.open(io.BytesIO(data))
            image.load()
            out = ASSET_DIR / f"image{i}.jpg"
            if i <= 17:
                rgb = image.convert("RGB")
                rgb.thumbnail((1400, 1400), Image.Resampling.LANCZOS)
                rgb.save(out, "JPEG", quality=86, optimize=True)
            else:
                rgb = image.convert("RGB")
                rgb.save(out, "JPEG", quality=94, optimize=True)


def add_picture_fit(paragraph, path, max_width_cm, max_height_cm):
    with Image.open(path) as image:
        width, height = image.size
    scale = min(max_width_cm / width, max_height_cm / height)
    return paragraph.add_run().add_picture(str(path), width=Cm(width * scale), height=Cm(height * scale))


def add_photo_grid(doc, image_numbers, captions, cols=3, box_w=5.35, box_h=8.2):
    rows = (len(image_numbers) + cols - 1) // cols
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color=WHITE, size=0)
    for idx in range(rows * cols):
        cell = table.cell(idx // cols, idx % cols)
        cell.width = Cm(5.65)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, top=65, start=60, bottom=80, end=60)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        if idx < len(image_numbers):
            num = image_numbers[idx]
            add_picture_fit(p, ASSET_DIR / f"image{num}.jpg", box_w, box_h)
            cp = cell.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_before = Pt(1)
            cp.paragraph_format.space_after = Pt(0)
            add_text(cp, captions[idx], size=8.2, color=GRAY)
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.55)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.paragraph_format.space_after = Pt(4)
        add_text(p, item, size=10.2)


def add_receipt_page(doc, title, description, images, layout="single"):
    doc.add_page_break()
    add_section_heading(doc, title, "증빙자료")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(9)
    add_text(p, description, size=10, color="444444")
    if layout == "side_by_side":
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_borders(table, color=WHITE, size=0)
        for cell, num in zip(table.rows[0].cells, images):
            cell.width = Cm(8.25)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=80, start=90, bottom=80, end=90)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_picture_fit(p, ASSET_DIR / f"image{num}.jpg", 7.7, 19.3)
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_picture_fit(p, ASSET_DIR / f"image{images[0]}.jpg", 16.0, 20.5)


extract_assets()
doc = Document()
configure_document(doc)
add_header_footer(doc)

add_title(doc)
add_info_table(doc)
add_callout(doc, "참가 목적", "학회 프로그램에 참여해 전공 분야의 연구 동향과 발표 사례를 살펴보고, 향후 연구 및 발표 준비에 참고할 내용을 정리하고자 했습니다.")
add_section_heading(doc, "1. 이동 및 일정")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
add_text(p, "8월 5일에는 한국에너지공과대학교에서 광주버스터미널로 이동한 뒤 경주로 출발했습니다. 8월 6일과 7일에는 현장 등록을 마치고 학회 일정에 참여했습니다.")
add_photo_grid(doc, [1, 2, 3], ["경주 이동 중 참가자 확인", "숙박 장소 외관", "숙박 시설 내부"], cols=3, box_w=5.2, box_h=6.3)

doc.add_page_break()
add_section_heading(doc, "2. 학회 프로그램 참석", "강연 및 발표 세션")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(7)
add_text(p, "학회장에서 전공 관련 강연과 Plenary Talk를 청취했습니다. 각 발표의 연구 배경, 접근 방법, 결과 제시 방식을 중심으로 내용을 살펴보았습니다.")
add_photo_grid(doc, [4, 5, 6, 7, 8, 9], ["강연 참석 1", "강연 참석 2", "강연 참석 3", "강연 참석 4", "강연 참석 5", "강연 참석 6"], cols=3, box_w=5.15, box_h=7.7)

doc.add_page_break()
add_section_heading(doc, "3. 현장 활동 기록", "강연장 및 포스터 세션")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(7)
add_text(p, "강연장과 포스터 세션을 오가며 여러 연구 주제와 발표 형식을 확인했습니다. 발표 자료의 구성과 핵심 내용을 전달하는 방식도 함께 살펴보았습니다.")
add_photo_grid(doc, [10, 11, 12, 13, 14, 15], ["강연 참석 7", "강연 참석 8", "강연 참석 9", "강연 참석 10", "포스터 세션", "강연 참석 11"], cols=3, box_w=5.15, box_h=7.7)

doc.add_page_break()
add_section_heading(doc, "4. 참가 결과 및 활용 계획")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(7)
add_text(p, "남은 프로그램에도 참석하며 학회 일정 전반을 확인했습니다. 이번 참가를 통해 연구 내용을 발표 자료로 정리하고 청중에게 전달하는 다양한 사례를 접할 수 있었습니다.")
add_photo_grid(doc, [16, 17], ["강연 참석 12", "강연 참석 13"], cols=2, box_w=5.1, box_h=7.8)
add_section_heading(doc, "주요 성과")
add_bullets(doc, [
    "강연과 포스터 발표를 통해 전공 분야에서 다뤄지는 연구 주제와 발표 흐름을 파악했습니다.",
    "연구 목적, 방법, 결과를 제한된 시간 안에 전달하는 발표 구성 방식을 확인했습니다.",
    "학회 현장 운영과 연구자 간 교류 분위기를 경험하며 향후 학술 활동 준비에 참고할 수 있는 사례를 수집했습니다.",
])
add_callout(doc, "활용 계획", "학회에서 확인한 발표 구성과 자료 제시 방식을 향후 연구 주제 검토, 발표 자료 작성 및 포스터 준비에 참고할 예정입니다.")

add_receipt_page(doc, "5. 학회 현장 등록 영수증", "학회 참가 등록과 관련한 증빙자료입니다.", [18])
add_receipt_page(doc, "6. 식비 영수증", "학회 참가 기간 중 발생한 식비 증빙자료입니다.", [19])
add_receipt_page(doc, "7. 교통비 영수증", "한국에너지공과대학교에서 경주로 이동하는 과정에서 발생한 교통비 증빙자료입니다.", [20, 21], layout="side_by_side")
add_receipt_page(doc, "8. 숙박비 영수증", "학회 참가 기간 중 이용한 숙박비 증빙자료입니다.", [22])

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
